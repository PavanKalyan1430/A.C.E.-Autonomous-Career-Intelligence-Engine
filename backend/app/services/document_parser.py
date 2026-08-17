import io
import asyncio
import logging
import os
from typing import Optional

logger = logging.getLogger(__name__)

# Magic-byte signatures for reliable file-type detection (RES-002).
# Extension-only checks are trivially bypassed by renaming files.
_MAGIC_BYTES: dict[str, bytes] = {
    ".pdf":  b"%PDF",
    ".docx": b"PK\x03\x04",  # DOCX/ZIP local file header
    ".doc":  b"\xd0\xcf\x11\xe0",  # OLE2 compound document (legacy .doc)
}


class DocumentParserService:
    """
    Production Multi-Format Document Ingestion Engine supporting PDF, DOCX, TXT, and Markdown files.
    Utilizes pypdf, python-docx, and unstructured partitioning with multi-tier fallback.

    All heavy synchronous parsing is offloaded to asyncio.to_thread (RES-007) to avoid
    blocking the FastAPI event loop on large documents.
    """

    async def extract_text(self, file_bytes: bytes, filename: str) -> str:
        """
        Entry point. Validates magic bytes then dispatches to the appropriate parser.
        Raises ValueError for unsupported or mismatched file types.
        """
        filename_lower = filename.lower()
        ext = os.path.splitext(filename_lower)[1]

        # --- Magic-byte validation (RES-002) ---
        self._validate_magic(file_bytes, ext, filename)

        # --- Dispatch to format-specific parser (RES-007: offload to thread) ---
        if ext == ".pdf":
            return await asyncio.to_thread(self._parse_pdf, file_bytes)
        elif ext in (".docx", ".doc"):
            return await asyncio.to_thread(self._parse_docx, file_bytes)
        else:
            # TXT / Markdown / unknown — plain-text decode (fast, no offload needed)
            return await asyncio.to_thread(self._parse_plain_text, file_bytes)

    def _validate_magic(self, file_bytes: bytes, ext: str, filename: str) -> None:
        """
        Cross-checks the first bytes of the file against the expected magic signature
        for the declared extension. Raises ValueError if they don't match.

        Only validates extensions that have known magic bytes; plain text has no
        universal magic signature and is skipped.
        """
        expected_magic = _MAGIC_BYTES.get(ext)
        if expected_magic is None:
            # .txt, .md, etc. — no magic to check
            return

        if not file_bytes[:len(expected_magic)] == expected_magic:
            raise ValueError(
                f"File content does not match declared extension '{ext}'. "
                f"File '{filename}' may be corrupted or misnamed."
            )

    def _parse_pdf(self, file_bytes: bytes) -> str:
        """Parses PDF documents using PyPDF and Unstructured with fallback."""
        text_chunks = []

        # 1. Try PyPDF stream extraction
        try:
            from pypdf import PdfReader
            pdf_stream = io.BytesIO(file_bytes)
            reader = PdfReader(pdf_stream)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text_chunks.append(extracted.strip())

            if text_chunks:
                logger.info(f"Successfully extracted {len(text_chunks)} PDF pages via PyPDF.")
                return "\n\n".join(text_chunks)
        except Exception as e:
            logger.warning(f"PyPDF extraction failed: {e}. Trying unstructured partition...")

        # 2. Try Unstructured partition fallback
        try:
            from unstructured.partition.pdf import partition_pdf
            pdf_stream = io.BytesIO(file_bytes)
            elements = partition_pdf(file=pdf_stream)
            extracted_text = "\n\n".join([str(el) for el in elements if str(el).strip()])
            if extracted_text:
                return extracted_text
        except Exception as e:
            logger.warning(f"Unstructured PDF partition failed: {e}.")

        # 3. Last-resort loose decoding
        return self._parse_plain_text(file_bytes)

    def _parse_docx(self, file_bytes: bytes) -> str:
        """Parses Microsoft Word (.docx) documents using python-docx and Unstructured."""
        text_chunks = []

        # 1. Try python-docx extraction
        try:
            import docx
            doc_stream = io.BytesIO(file_bytes)
            document = docx.Document(doc_stream)

            # Extract paragraphs
            for para in document.paragraphs:
                if para.text.strip():
                    text_chunks.append(para.text.strip())

            # Extract tables
            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_chunks.append(row_text)

            if text_chunks:
                logger.info(f"Successfully extracted DOCX document via python-docx.")
                return "\n\n".join(text_chunks)
        except Exception as e:
            logger.warning(f"python-docx extraction failed: {e}. Trying unstructured partition...")

        # 2. Try Unstructured docx partition
        try:
            from unstructured.partition.docx import partition_docx
            doc_stream = io.BytesIO(file_bytes)
            elements = partition_docx(file=doc_stream)
            extracted_text = "\n\n".join([str(el) for el in elements if str(el).strip()])
            if extracted_text:
                return extracted_text
        except Exception as e:
            logger.warning(f"Unstructured DOCX partition failed: {e}.")

        return self._parse_plain_text(file_bytes)

    def _parse_plain_text(self, file_bytes: bytes) -> str:
        """Decodes plain text/markdown bytes with multi-encoding fallback."""
        for encoding in ["utf-8", "utf-16", "latin-1"]:
            try:
                return file_bytes.decode(encoding).strip()
            except UnicodeDecodeError:
                continue
        return file_bytes.decode("utf-8", errors="ignore").strip()


document_parser_service = DocumentParserService()
